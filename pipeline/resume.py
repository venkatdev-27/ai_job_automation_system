import json
import re
import os
import sys
import hashlib
import requests
import asyncio
from pathlib import Path
from datetime import datetime
from .config import TEMP_DIR, C, LINKEDIN_EMAIL, GROQ_MODEL, GROQ_API_KEY, GROQ_API_BASE
from .pdf_service import pdf_service
from .llm import groq_call, parse_json_from_llm

def log_info(msg): print(f"{C.CYAN}[INFO] {msg}{C.RESET}")
def log_ok(msg):   print(f"{C.GREEN}[OK] {msg}{C.RESET}")
def log_warn(msg): print(f"{C.YELLOW}[WARN] {msg}{C.RESET}")
def log_err(msg):  print(f"{C.RED}[ERR] {msg}{C.RESET}")

# --- Step 1: Download & Text Extraction ---

def download_cloudinary_resume(url: str) -> Path:
    log_info(f"Downloading resume from Cloudinary: {url}")
    url_hash = hashlib.md5(url.encode()).hexdigest()[:8]
    dest = TEMP_DIR / f"source_resume_{url_hash}.docx"
    if dest.exists():
        return dest
    r = requests.get(url, timeout=30)
    r.raise_for_status()
    dest.write_bytes(r.content)
    log_ok(f"Downloaded resume to {dest}")
    return dest

def extract_docx_text(docx_path: Path) -> str:
    from docx import Document
    doc = Document(str(docx_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs)

def extract_skills_from_text(text: str) -> list:
    TECH_SKILLS = [
        "react", "node.js", "nodejs", "express", "mongodb", "javascript", "typescript",
        "html", "css", "python", "java", "spring", "angular", "vue", "docker",
        "kubernetes", "aws", "azure", "git", "redux", "rest api", "graphql",
        "mysql", "postgresql", "redis", "next.js", "nestjs", "tailwind",
        "jest", "mocha", "webpack", "ci/cd", "linux", "agile", "scrum",
        "c#", ".net", "asp.net", "sql", "firebase", "flutter", "react native"
    ]
    text_lower = text.lower()
    found = []
    for skill in TECH_SKILLS:
        if skill in text_lower:
            found.append(skill)
    return found[:20]

def detect_role_from_resume(resume_text: str) -> dict:
    log_info("Detecting role and profile from resume using Groq LLM...")
    prompt = f"""Analyze this resume and extract the candidate profile as JSON.
RESUME TEXT:
{resume_text[:4000]}
Return ONLY this JSON:
{{
  "full_name": "...",
  "email": "...",
  "phone": "...",
  "location": "...",
  "linkedin": "...",
  "github": "...",
  "portfolio": "...",
  "target_role": "...",
  "search_keywords": ["keyword1", "keyword2"],
  "skills": {{
    "Dynamic Category 1": ["skill1", "skill2"],
    "Dynamic Category 2": ["..."]
  }},
  "experience": [
    {{
      "role": "Title",
      "company": "Company",
      "duration": "Year-Year",
      "bullets": ["Bullet 1", "Bullet 2"]
    }}
  ],
  "projects": [
    {{
      "name": "Project Name",
      "tech": "Stack",
      "bullets": ["Detail 1", "Detail 2"]
    }}
  ],
  "education": "Degree at Institution, Year",
  "summary": "..."
}}"""
    raw = groq_call(prompt)
    profile = parse_json_from_llm(raw)
    if not profile.get("target_role"):
        profile["target_role"] = "Software Developer"
    return profile

# --- Step 2: Template Analysis ---

def extract_resume_template_profile(docx_path: Path, resume_text: str) -> dict:
    profile = {
        "font_family": "Calibri",
        "body_font_size_pt": 11,
        "accent_color": "#2C5282",
        "skill_categories": [],
    }
    # Simplified extraction for modular version
    profile["skill_categories"] = ["Frontend", "Backend", "Database", "Tools & DevOps"]
    return profile

# --- Step 3: Tailoring & Validation ---

def validate_resume_truth(original_text: str, generated_data: dict) -> list:
    """Check if the generated resume data contains hallucinations not in the original text."""
    log_info("Running truth validation check...")
    
    # We send a summary of the facts and the new content to the LLM to verify
    prompt = f"""
I am auditing a generated resume for truthfulness. 
ORIGINAL RESUME TEXT:
{original_text[:4000]}

GENERATED RESUME JSON (SUMMARY):
Summary: {generated_data.get('summary')}
Experience: {json.dumps(generated_data.get('experience', []))[:2000]}
Skills: {json.dumps(generated_data.get('skills', {}))}

Identify any specific skills, technologies, companies, or major achievements in the GENERATED version that are NOT present (directly or strongly implied) in the ORIGINAL version.
Return a JSON list of strings (hallucinations found). If none, return empty list [].
Return only JSON.
"""
    raw = groq_call(prompt, system="You are a strict technical auditor.")
    hallucinations = parse_json_from_llm(raw)
    if isinstance(hallucinations, list) and hallucinations:
        log_warn(f"Hallucination check flagged {len(hallucinations)} items: {hallucinations}")
    else:
        log_ok("Truth validation passed.")
    return hallucinations or []

def generate_tailored_resume_optimized(resume_text: str, job_description: str, profile: dict, template_profile: dict, job_points: list = None) -> dict:
    """Combined Tailoring and Polishing in a single LLM pass."""
    log_info("Generating tailored + polished resume (Single Pass optimization)...")
    
    points_str = f"\nKEY JOB REQUIREMENTS:\n- " + "\n- ".join(job_points) if job_points else ""
    
    system = """You are an expert ATS resume writer. 
Rules:
1. ONLY use experience and skills actually present in the original resume.
2. Optimization: Use powerful action verbs and technical specifics. 
3. Match keywords from the job description if the candidate has those skills.
4. Output must be perfectly valid JSON."""

    prompt = f"""
ORIGINAL RESUME:
{resume_text[:3500]}

TARGET JOB DESCRIPTION (600 Words):
{job_description[:4500]}
{points_str}

Generate a complete resume JSON tailored to this JD.
Structure:
{{
  "full_name": "{profile.get('full_name', 'Candidate')}",
  "target_role": "...",
  "email": "{profile.get('email', '')}",
  "phone": "{profile.get('phone', '')}",
  "location": "{profile.get('location', '')}",
  "summary": "3-4 concise sentences...",
  "skills": {{ "Category": ["Skill 1", "Skill 2"] }},
  "experience": [ {{ "title": "...", "company": "...", "duration": "...", "bullets": ["Result-oriented bullet 1..."] }} ],
  "projects": [ {{ "name": "...", "tech": "...", "bullets": ["..."] }} ],
  "education": {{ "degree": "...", "institution": "...", "year": "..." }},
  "ats_keywords_used": ["k1", "k2"]
}}
"""
    raw = groq_call(prompt, system=system)
    data = parse_json_from_llm(raw)
    
    # Truth validation step
    hallucinations = validate_resume_truth(resume_text, data)
    if hallucinations:
        # Subtle self-correction or just log for now
        data["notes"] = f"Warning: Potential hallucinations flagged: {hallucinations}"
        
    return data

# --- Step 4: Output Generation ---

def generate_resume_html(resume_data: dict, template_profile: dict | None = None) -> str:
    log_info("Generating ATS-friendly HTML resume...")
    template_profile = template_profile or {}
    
    name = resume_data.get("full_name", "Candidate")
    role = resume_data.get("target_role", "Software Developer")
    email = resume_data.get("email", "")
    phone = resume_data.get("phone", "")
    location = resume_data.get("location", "")
    linkedin = resume_data.get("linkedin", "")
    github = resume_data.get("github", "")
    summary = resume_data.get("summary", "")

    font_family = template_profile.get("font_family", '"Times New Roman", Times, serif')
    body_font_size = template_profile.get("body_font_size_pt", 11)
    name_font_size = template_profile.get("name_font_size_pt", max(18, body_font_size + 10))
    section_title_size = template_profile.get("section_title_size_pt", body_font_size + 1)
    summary_size = template_profile.get("summary_size_pt", max(10, body_font_size - 0.5))
    accent_color = template_profile.get("accent_color", "#2C5282")
    primary_color = template_profile.get("primary_color", "#1A365D")
    contact_color = template_profile.get("contact_color", "#555555")
    header_alignment = template_profile.get("header_alignment", "center")
    section_order = template_profile.get(
        "section_order",
        ["Professional Summary", "Technical Skills", "Work Experience", "Projects", "Education", "Certifications"],
    )
    
    # Build skills HTML
    skills_html = ""
    skills = resume_data.get("skills", {})
    if isinstance(skills, dict):
        for cat, skill_list in skills.items():
            if skill_list:
                skills_html += f'<tr><td class="skill-cat"><b>{cat}:</b></td><td class="skill-items">{", ".join(skill_list)}</td></tr>'
    elif isinstance(skills, list):
        skills_html = f'<tr><td class="skill-cat"><b>Skills:</b></td><td class="skill-items">{", ".join(skills)}</td></tr>'

    # Build experience HTML
    exp_html = ""
    for exp in resume_data.get("experience", []):
        bullets = "".join([f"<li>{b}</li>" for b in exp.get("bullets", [])])
        exp_html += f"""
        <div class="exp-entry">
          <div class="exp-header">
            <span class="exp-title">{exp.get('title', '')}</span>
            <span class="exp-company">{exp.get('company', '')}</span>
            <span class="exp-duration">{exp.get('duration', '')}</span>
          </div>
          <ul class="bullets">{bullets}</ul>
        </div>"""

    # Build projects HTML
    proj_html = ""
    for proj in resume_data.get("projects", []):
        bullets = "".join([f"<li>{b}</li>" for b in proj.get("bullets", [])])
        proj_html += f"""
        <div class="exp-entry">
          <div class="exp-header">
            <span class="exp-title">{proj.get('name', '')}</span>
            <span class="exp-company">{proj.get('tech', '')}</span>
          </div>
          <ul class="bullets">{bullets}</ul>
        </div>"""

    # Education
    edu = resume_data.get("education", {})
    edu_html = ""
    if isinstance(edu, dict):
        edu_html = f'<p><b>{edu.get("degree", "")}</b> — {edu.get("institution", "")} {edu.get("year", "")}</p>'
    elif isinstance(edu, str):
        edu_html = f"<p>{edu}</p>"

    # Certifications
    certs = resume_data.get("certifications", [])
    certs_html = ""
    if certs:
        certs_html = f'<section class="section"><h2 class="section-title">Certifications</h2><ul class="bullets">{"".join([f"<li>{c}</li>" for c in certs])}</ul></section>'

    # Contact line
    contact_parts = [p for p in [email, phone, location, linkedin, github] if p]
    contact_line = " | ".join(contact_parts)

    sections = {
        "Professional Summary": f"""
  <section class="section">
    <h2 class="section-title">Professional Summary</h2>
    <p class="summary-text">{summary}</p>
  </section>""",
        "Technical Skills": f"""
  <section class="section">
    <h2 class="section-title">Technical Skills</h2>
    <table class="skills-table">
      {skills_html}
    </table>
  </section>""",
        "Work Experience": "<section class='section'><h2 class='section-title'>Work Experience</h2>" + exp_html + "</section>" if exp_html else "",
        "Projects": "<section class='section'><h2 class='section-title'>Projects</h2>" + proj_html + "</section>" if proj_html else "",
        "Education": f"""
  <section class="section">
    <h2 class="section-title">Education</h2>
    <div class="edu-content">{edu_html}</div>
  </section>""",
        "Certifications": certs_html,
    }

    resolved_order = []
    for section in section_order:
        canonical = str(section).strip()
        if canonical in sections and canonical not in resolved_order:
            resolved_order.append(canonical)
    for fallback in sections.keys():
        if fallback not in resolved_order:
            resolved_order.append(fallback)
    body_sections = "\n".join([sections[s] for s in resolved_order if sections[s]])

    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{name} - Resume</title>
<style>
  * {{ margin: 0; padding: 0; box-sizing: border-box; }}
  body {{ font-family: {font_family}, Arial, sans-serif; font-size: {body_font_size}pt; line-height: 1.4; color: #222; padding: 0.6in; max-width: 8.5in; }}
  .header {{ text-align: {header_alignment}; border-bottom: 2px solid {accent_color}; padding-bottom: 10px; margin-bottom: 12px; }}
  .name {{ font-size: {name_font_size}pt; font-weight: bold; color: {primary_color}; letter-spacing: 0.5px; }}
  .role-title {{ font-size: {section_title_size}pt; color: {accent_color}; margin: 4px 0; }}
  .contact {{ font-size: {max(9, body_font_size - 1.5)}pt; color: {contact_color}; margin-top: 4px; }}
  .section {{ margin-bottom: 14px; }}
  .section-title {{ font-size: {section_title_size}pt; font-weight: bold; color: {primary_color}; text-transform: uppercase; letter-spacing: 0.8px; border-bottom: 1px solid #a0aec0; padding-bottom: 3px; margin-bottom: 8px; }}
  .summary-text {{ font-size: {summary_size}pt; color: #333; line-height: 1.5; text-align: justify; }}
  .skills-table {{ width: 100%; border-collapse: collapse; font-size: {max(9.5, body_font_size - 1)}pt; }}
  .skill-cat {{ width: 22%; padding: 2px 8px 2px 0; vertical-align: top; color: {accent_color}; white-space: nowrap; }}
  .skill-items {{ padding: 2px 0; }}
  .exp-entry {{ margin-bottom: 10px; }}
  .exp-header {{ display: flex; flex-wrap: wrap; gap: 6px; margin-bottom: 4px; }}
  .exp-title {{ font-weight: bold; color: {primary_color}; font-size: {summary_size}pt; }}
  .exp-company {{ color: #555; font-style: italic; }}
  .exp-duration {{ margin-left: auto; color: #777; font-size: {max(9, body_font_size - 1.5)}pt; }}
  .bullets {{ margin-left: 18px; font-size: {max(9.5, body_font_size - 1)}pt; }}
  .bullets li {{ margin-bottom: 3px; line-height: 1.4; }}
  .edu-content {{ font-size: {summary_size}pt; }}
</style>
</head>
<body>
  <div class="header">
    <div class="name">{name}</div>
    <div class="role-title">{role}</div>
    <div class="contact">{contact_line}</div>
  </div>

  {body_sections}
</body>
</html>"""
    return html

def check_ats_score(resume_data: dict, job_description: str, job_skills: list) -> dict:
    log_info("Running ATS score analysis...")
    
    # Flatten resume skills
    resume_skills_raw = resume_data.get("skills", {})
    if isinstance(resume_skills_raw, dict):
        all_resume_skills = []
        for v in resume_skills_raw.values():
            all_resume_skills.extend([s.lower() for s in v])
    else:
        all_resume_skills = [s.lower() for s in resume_skills_raw]

    jd_lower = job_description.lower()
    
    # Keyword matching
    matched = [s for s in all_resume_skills if s in jd_lower]
    keyword_score = min(100, int((len(matched) / max(len(job_skills), 1)) * 100))
    
    return {"ats_score": keyword_score, "matched_keywords": matched}

async def html_to_pdf(html_content: str, output_path: Path):
    await pdf_service.generate_pdf(html_content, output_path)
